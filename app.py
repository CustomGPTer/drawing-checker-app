from flask import Flask, request, render_template, jsonify, send_file
import os, zipfile, fitz, ezdxf, shutil, uuid, tempfile, re
from datetime import datetime, timedelta
from werkzeug.utils import secure_filename
from openai import OpenAI
from docx import Document

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
PROCESSED_FOLDER = 'processed_reports'
REFERENCE_DRAWINGS_DIR = 'reference_drawings_extracted'
REFERENCE_ZIP = 'reference_drawings/master_drawings.zip'
REFERENCE_DOCS_FOLDER = 'reference_docs'

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)
os.makedirs(REFERENCE_DRAWINGS_DIR, exist_ok=True)

reference_specs = {}
reference_drawings_text = {}
def read_pdf_text(path):
    doc = fitz.open(path)
    return "\n".join(page.get_text() for page in doc)

def read_dxf_text(path):
    doc = ezdxf.readfile(path)
    return "\n".join([e.dxf.text for e in doc.modelspace() if hasattr(e.dxf, "text")])

def extract_zip(zip_path, extract_to):
    with zipfile.ZipFile(zip_path, 'r') as zip_ref:
        zip_ref.extractall(extract_to)

def load_reference_specs():
    for filename in os.listdir(REFERENCE_DOCS_FOLDER):
        path = os.path.join(REFERENCE_DOCS_FOLDER, filename)
        if filename.endswith('.pdf'):
            try:
                text = read_pdf_text(path)
                reference_specs[filename] = text[:10000]
            except Exception as e:
                reference_specs[filename] = f"⚠️ Could not read: {e}"

if os.path.exists(REFERENCE_ZIP):
    extract_zip(REFERENCE_ZIP, REFERENCE_DRAWINGS_DIR)
    for file in os.listdir(REFERENCE_DRAWINGS_DIR):
        path = os.path.join(REFERENCE_DRAWINGS_DIR, file)
        if file.endswith('.pdf'):
            reference_drawings_text[file] = read_pdf_text(path)
        elif file.endswith(('.dxf', '.dwg')):
            try:
                reference_drawings_text[file] = read_dxf_text(path)
            except:
                reference_drawings_text[file] = 'Unreadable DWG'

load_reference_specs()

def cleanup_old_sessions():
    now = datetime.now()
    for folder in os.listdir(UPLOAD_FOLDER):
        folder_path = os.path.join(UPLOAD_FOLDER, folder)
        if os.path.isdir(folder_path):
            try:
                folder_time = datetime.strptime(folder.split('_')[0], '%Y%m%d%H%M%S')
                if now - folder_time > timedelta(days=7):
                    shutil.rmtree(folder_path)
            except:
                pass

def get_latest_revisions(folder):
    drawings = {}
    for filename in os.listdir(folder):
        if not filename.lower().endswith(('.pdf', '.dwg', '.dxf')):
            continue
        match = re.match(r'(DR-[A-Z]+-\d+)-([CPD]\d+)\.', filename)
        if match:
            base = match.group(1)
            rev = match.group(2)
            current = drawings.get(base)
            if not current or rev[1:] > current[1][1:]:
                drawings[base] = (filename, rev)
        else:
            # Include non-matching files using filename as fallback
            drawings[filename] = (filename, "NA")
    return [os.path.join(folder, v[0]) for v in drawings.values()]

    
openai = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

def generate_prompt(drawing_number, title, revision, text, reference_texts, ref_drawings):
    spec_context = "\n".join([f"{name}:\n{doc[:1000]}" for name, doc in reference_texts.items()])
    zip_context = "\n".join([f"{name}:\n{text[:1000]}" for name, text in ref_drawings.items()])
    
    qa_checks = '''
1. ✅ Are all cover/invert levels shown, consistent, and buildable?
2. ✅ Are pipe bedding types correct per CESWI/UUCESWI?
3. ✅ Is flow direction clearly shown?
4. ✅ Do chamber references and layouts match schedules?
5. ✅ Are wall, slab, and foundation thicknesses shown and labelled?
6. ✅ Is reinforcement correctly detailed?
7. ✅ Are pipe sizes, gradients, and materials labelled?
8. ✅ Are access ladders, platforms, or landings included where needed?
9. ✅ Are chambers and covers accessible for lifting and maintenance?
10. ✅ Are plans, sections, and detail views coordinated?
11. ✅ Are cable tray routes clash-free with civils?
12. ✅ Are ducts shown with correct layout, spacing, and annotation?
13. ✅ Are drawpits and duct bends buildable and spaced to spec?
14. ✅ Are pumps and valves fully detailed and accessible?
15. ✅ Are mechanical/electrical isolations clearly marked?
16. ✅ Are sensors and instruments located correctly?
17. ✅ Are control panels coordinated with structure?
18. ✅ Are civils penetrations shown for M&E systems?
19. ✅ Are vent routes logical and clash-free?
20. ✅ Are bonding and earthing points compliant?
21. ✅ Does this drawing comply with CESWI/UUCESWI?
22. ✅ Are United Utilities (UU) standard details applied correctly?
23. ✅ Is this the current approved-for-construction revision?
24. ✅ Are referenced drawings accurate and consistent?
25. ✅ Are temporary works or staged build notes included where required?
26. ✅ Are maintenance/lifting zones and fall protection shown?
27. ✅ Is the scope clear in the title block and general notes?
28. ✅ Is the drawing coordinated with current RAMS and construction methods?
29. ✅ Are services and structures shown logically and buildable?
30. ✅ Are there any omissions identified using your internal engineering knowledge?
'''

    return f'''
You are a construction drawing checker built for C2V+ projects working on United Utilities infrastructure sites.

A ZIP file of reference drawings has been uploaded. You must fully read and cross-reference all files, then assess the uploaded drawing.

Use CESWI 7th Edition, UUCESWI amendments, and C2V+ "What Good Looks Like" standards.

### Drawing Details:
- Drawing Number: {drawing_number}
- Title: {title}
- Revision: {revision}

--- Drawing Content ---
{text}

--- Reference Documents ---
{spec_context}

--- Master Drawings ---
{zip_context}

### Instructions:
1. Identify the drawing type (e.g., Drainage Layout, Cable Routing, Valve Chamber)
2. Apply all relevant checks from the following 30-point QA list:
{qa_checks}

3. Output results in this format:

---
Result: ✅ / ⚠️ / ❌  
Explanation (technical, specific)  
Drawing Reference  
Suggested Action  
---

Repeat for each check.

Score the drawing out of 30.  
Then give:
- Risk Level: Low / Medium / High
- Additional Observations
- Clashes or omissions across documents
- Notes for further clarification (e.g. request plan/section views)

Only refer to visible content — never assume.
'''


def call_gpt(prompt):
    response = openai.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2
    )
    return response.choices[0].message.content

def score_compliance(report_text):
    total = 0.0
    for line in report_text.splitlines():
        if line.strip().startswith("Result:"):
            if '✅' in line: total += 1
            elif '⚠️' in line: total += 0.5
    if total >= 27: risk = "Low"
    elif total >= 20: risk = "Medium"
    else: risk = "High"
    return total, risk

def annotate_pdf(input_path, comments):
    doc = fitz.open(input_path)
    for page in doc:
        for text in comments:
            if text.lower() in page.get_text().lower():
                page.insert_text((50, 50 + 20*comments.index(text)), f"❌ {text[:80]}", fontsize=8, color=(1, 0, 0))
    out_path = input_path.replace('.pdf', '_annotated.pdf')
    doc.save(out_path)
    return out_path

def generate_docx_report(drawing_name, results, score, risk):
    doc = Document()
    doc.add_heading(f"Drawing QA Report: {drawing_name}", level=1)
    doc.add_paragraph(f"Compliance Score: {score}/30  –  Risk Level: {risk}")
    for section in results.split('\n\n'):
        doc.add_paragraph(section)
    filename = f"{drawing_name.replace('.', '_')}_QA_Report.docx"
    path = os.path.join(PROCESSED_FOLDER, filename)
    doc.save(path)
    return path
@app.route('/')
def index():
    return render_template('upload.html')

@app.route('/upload', methods=['POST'])
def upload():
    cleanup_old_sessions()
    files = request.files.getlist('drawings')
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    session_id = f"{timestamp}_{str(uuid.uuid4())[:8]}"
    session_folder = os.path.join(UPLOAD_FOLDER, session_id)
    os.makedirs(session_folder)

    zip_files = []
    user_files = []

    for file in files:
        filename = secure_filename(file.filename)
        path = os.path.join(session_folder, filename)
        file.save(path)
        if filename.endswith('.zip'):
            zip_files.append(path)
        else:
            user_files.append(path)

    for zip_path in zip_files:
        extract_zip(zip_path, session_folder)

    all_drawings = get_latest_revisions(session_folder)
    drawing_index = {}
    for path in all_drawings:
        filename = os.path.basename(path)
        match = re.match(r'(DR-[A-Z]+-\d+)-([CPD]\d+)\.', filename)
        number = match.group(1) if match else filename
        drawing_index.setdefault(number, []).append(filename)

    with open(os.path.join(session_folder, 'drawing_index.json'), 'w') as f:
        f.write(str(drawing_index))

    report_links = []
    summary_table = []

    for path in all_drawings:
        filename = os.path.basename(path)
        drawing_number = filename.split('-')[2] if '-' in filename else filename
        title = filename.rsplit('.', 1)[0]
        revision = filename.split('-')[-1].split('.')[0]

        try:
            if filename.endswith('.pdf'):
                text = read_pdf_text(path)
            elif filename.endswith(('.dwg', '.dxf')):
                text = read_dxf_text(path)
            else:
                continue

            prompt = generate_prompt(drawing_number, title, revision, text, reference_specs, reference_drawings_text)
            gpt_result = call_gpt(prompt)
            score, risk = score_compliance(gpt_result)
            docx_path = generate_docx_report(filename, gpt_result, score, risk)
            report_links.append({"docx": docx_path})

            if filename.endswith('.pdf'):
                flagged_notes = [line for line in gpt_result.splitlines() if '❌' in line or '⚠️' in line]
                overlay_path = annotate_pdf(path, flagged_notes)
                report_links[-1]['pdf_overlay'] = overlay_path

            summary_table.append({
                "drawing": filename,
                "score": score,
                "risk": risk
            })

        except Exception as e:
            report_links.append({"error": f"Failed to process {filename}: {e}"})

    return jsonify({
        "session_id": session_id,
        "drawing_index": drawing_index,
        "reports": report_links,
        "summary": summary_table
    })
if __name__ == "__main__":
    app.run(debug=True)
